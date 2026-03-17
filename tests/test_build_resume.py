# tests/test_build_resume.py
import pytest
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "tools"))
from build_resume import validate


VALID = {
    "name": "Kyaw Min Thu", "email": "kyaw@kmtkn.me",
    "phone": "+1 510 909 3716", "location": "Bay Area, CA",
    "company": "Acme", "role": "Engineer",
    "summary": "A summary.",
    "competencies": ["PLC", "SCADA"],
    "experience": [{"company": "Co", "role": "Eng", "dates": "2020–2024",
                    "bullets": ["Did thing."]}],
    "education": [{"degree": "BS CS", "institution": "CSUEB",
                   "location": "Hayward, CA", "dates": "2018–2024"}],
}


def test_valid_data_passes():
    validate(VALID)  # should not raise


def test_missing_name_raises():
    with pytest.raises(ValueError, match="name"):
        validate({**VALID, "name": None})


def test_missing_email_raises():
    with pytest.raises(ValueError, match="email"):
        validate({**VALID, "email": None})


def test_missing_company_raises():
    with pytest.raises(ValueError, match="company"):
        validate({**VALID, "company": None})


def test_missing_role_raises():
    with pytest.raises(ValueError, match="role"):
        validate({**VALID, "role": None})


def test_missing_phone_raises():
    with pytest.raises(ValueError, match="phone"):
        validate({**VALID, "phone": None})


def test_missing_location_raises():
    with pytest.raises(ValueError, match="location"):
        validate({**VALID, "location": None})


def test_missing_summary_raises():
    with pytest.raises(ValueError, match="summary"):
        validate({**VALID, "summary": None})


def test_missing_competencies_raises():
    with pytest.raises(ValueError, match="competencies"):
        validate({**VALID, "competencies": None})


def test_missing_education_raises():
    with pytest.raises(ValueError, match="education"):
        validate({**VALID, "education": None})


def test_empty_experience_raises():
    with pytest.raises(ValueError, match="experience"):
        validate({**VALID, "experience": []})


def test_missing_bullets_raises():
    bad_exp = [{"company": "Co", "role": "Eng", "dates": "2020"}]
    with pytest.raises(ValueError, match="bullets"):
        validate({**VALID, "experience": bad_exp})


from build_resume import sanitize_filename


def test_spaces_become_underscores():
    assert sanitize_filename("Amazon Robotics") == "Amazon_Robotics"


def test_special_chars_stripped():
    assert sanitize_filename("C&W / JLL") == "CW__JLL"


def test_capped_at_50_chars():
    long = "A" * 60
    assert len(sanitize_filename(long)) == 50


def test_empty_string_returns_unknown():
    assert sanitize_filename("") == "Unknown"


def test_none_returns_unknown():
    assert sanitize_filename(None) == "Unknown"


import build_resume
from docx.shared import Pt, Inches
from build_resume import build_header, make_doc


def test_make_doc_margins():
    doc = make_doc()
    section = doc.sections[0]
    assert section.left_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert section.top_margin == Inches(1)
    assert section.bottom_margin == Inches(1)


def test_make_doc_no_default_paragraph():
    doc = make_doc()
    assert len(doc.paragraphs) == 0


def test_header_name_is_first_paragraph():
    doc = make_doc()
    build_header(doc, VALID)
    assert doc.paragraphs[0].text == "Kyaw Min Thu"


def test_header_name_is_bold():
    doc = make_doc()
    build_header(doc, VALID)
    assert doc.paragraphs[0].runs[0].bold is True


def test_header_name_font_size():
    doc = make_doc()
    build_header(doc, VALID)
    assert doc.paragraphs[0].runs[0].font.size == Pt(16)


def test_header_name_font_is_calibri():
    doc = make_doc()
    build_header(doc, VALID)
    assert doc.paragraphs[0].runs[0].font.name == "Calibri"


def test_header_contact_line_contains_email():
    doc = make_doc()
    build_header(doc, VALID)
    assert "kyaw@kmtkn.me" in doc.paragraphs[1].text


def test_header_contact_line_contains_phone():
    doc = make_doc()
    build_header(doc, VALID)
    assert "+1 510 909 3716" in doc.paragraphs[1].text
