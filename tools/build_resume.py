import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


REQUIRED_TOP_LEVEL = ["name", "email", "phone", "location", "summary",
                       "competencies", "experience", "education", "company", "role"]


def validate(data: dict) -> None:
    """Raise ValueError with field name if any required field is missing or null."""
    for field in REQUIRED_TOP_LEVEL:
        if data.get(field) is None:
            raise ValueError(f"Missing required field: '{field}'")
    if not data["experience"]:
        raise ValueError("'experience' list is empty")
    for i, job in enumerate(data["experience"]):
        if "bullets" not in job:
            raise ValueError(f"experience[{i}] missing 'bullets' key")


def sanitize_filename(value: str | None) -> str:
    if not value:
        return "Unknown"
    cleaned = re.sub(r'[^A-Za-z0-9_-]', '', value.replace(' ', '_'))
    return cleaned[:50] if cleaned else "Unknown"


def make_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    # Remove default empty paragraph
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)
    return doc


def _set_font(run, size_pt: int, bold: bool = False):
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.bold = bold


def build_header(doc: Document, data: dict) -> None:
    # Name
    p = doc.add_paragraph()
    run = p.add_run(data["name"])
    _set_font(run, 16, bold=True)

    # Contact line
    contact = f"{data['email']}  ·  {data['phone']}  ·  {data['location']}"
    p2 = doc.add_paragraph()
    run2 = p2.add_run(contact)
    _set_font(run2, 10)
